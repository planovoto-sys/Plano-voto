from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "generated-app-docs"
DOCX_PATH = OUT_DIR / "documentacao-estrutura-diretorios-nossovoto.docx"

ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 101, 119)
TABLE_HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F7FAFC"
TABLE_BORDER = "CBD5E1"


def text_file_line_count(path: Path) -> str:
    if not path.exists() or path.is_dir():
        return "-"
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".ico", ".pdf", ".docx"}:
        return "-"
    if path.name == ".env.local":
        return "privado"
    try:
        return str(len(path.read_text(encoding="utf-8", errors="ignore").splitlines()))
    except OSError:
        return "-"


def file_size(path: Path) -> str:
    if not path.exists():
        return "-"
    size = path.stat().st_size
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 * 1024):.1f} MB"


def rel(path: str) -> str:
    return path.replace("/", "\\")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), TABLE_BORDER)


def set_table_grid(table, widths) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in [
        ("Title", 24, RGBColor(17, 24, 39), 0, 7),
        ("Subtitle", 10.5, MUTED, 0, 10),
        ("Heading 1", 15, ACCENT, 10, 5),
        ("Heading 2", 12.3, ACCENT, 8, 4),
        ("Heading 3", 10.8, ACCENT_DARK, 6, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if style_name.startswith("Heading"):
            style.font.bold = True


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.text = "nossovoto.org | Organizacao de pastas e arquivos"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

    footer_para = section.footer.paragraphs[0]
    footer_para.text = f"Gerado em {datetime.now().strftime('%d/%m/%Y')} | Repositorio local Plano-voto"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def add_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    set_table_grid(table, [9300])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = ACCENT_DARK
    cell.add_paragraph(body)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = hdr_cells[idx]
        set_cell_shading(cell, TABLE_HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 41, 55)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(8.3)
    set_table_borders(table)
    set_table_grid(table, widths)
    doc.add_paragraph()


def add_file_table(doc: Document, rows: list[dict[str, str]], widths: list[int] | None = None) -> None:
    rendered_rows = []
    for row in rows:
        path = row["path"]
        path_obj = ROOT / path
        rendered_rows.append([
            rel(path),
            row["area"],
            file_size(path_obj),
            text_file_line_count(path_obj),
            row["purpose"],
        ])
    add_table(
        doc,
        ["Arquivo", "Area", "Tam.", "Linhas", "O que faz"],
        rendered_rows,
        widths or [2450, 1350, 780, 720, 4000],
    )


def add_dir_table(doc: Document, rows: list[list[str]]) -> None:
    add_table(doc, ["Pasta", "Classificacao", "Organizacao / responsabilidade"], rows, [2200, 1900, 5200])


def build_tree_text() -> str:
    return "\n".join([
        "Plano-voto/",
        "  src/                         Front-end React e logicas de UI",
        "    app/                       Roteamento principal",
        "    assets/                    Imagens e marca usadas pelo front",
        "    components/                Pecas reutilizaveis de interface",
        "    constants/                 Rotas, estados, cargos, textos legais",
        "    contexts/, hooks/, providers/",
        "                               Estado global de usuario e hooks",
        "    pages/                     Telas completas da aplicacao",
        "    services/                  Camada cliente: Firebase, cache, votos, share, PWA",
        "    styles/                    CSS global, reset e Tailwind v4",
        "    utils/                     Funcoes puras de apoio",
        "  public/                      Arquivos estaticos servidos pela PWA",
        "  functions/                   Back-end serverless Firebase Functions",
        "  docs/                        Documentacao e materiais gerados",
        "  dist/                        Build gerado pelo Vite",
        "  node_modules/                Dependencias instaladas",
        "  *.json, *.js, *.md, .env*    Configuracao, deploy e metadados",
    ])


def add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(0.2)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(7.9)
        run.font.color.rgb = RGBColor(55, 65, 81)


ROOT_FILES = [
    {
        "path": ".env.example",
        "area": "Config",
        "purpose": "Modelo das variaveis VITE_* exigidas pelo front-end e nomes das Cloud Functions. Deve ser copiado para ambiente local sem valores sensiveis.",
    },
    {
        "path": ".env.local",
        "area": "Config privado",
        "purpose": "Arquivo local ignorado pelo Git com valores reais de Firebase/Vite. Nao deve ser compartilhado nem copiado para documentacao.",
    },
    {
        "path": ".gitignore",
        "area": "Git",
        "purpose": "Define artefatos que nao devem entrar no repositorio: logs, node_modules, dist, arquivos .local e pastas de editor.",
    },
    {
        "path": "CHANGELOG.md",
        "area": "Docs",
        "purpose": "Registro historico de mudancas do projeto. Ajuda a entender evolucao de versoes e entregas.",
    },
    {
        "path": "README.md",
        "area": "Docs",
        "purpose": "Apresentacao do projeto, tecnologias, funcionalidades, screenshots e estrutura resumida.",
    },
    {
        "path": "package.json",
        "area": "Build front",
        "purpose": "Manifesto npm principal. Define app Vite/React, scripts dev/build/lint/deploy e dependencias de front-end.",
    },
    {
        "path": "package-lock.json",
        "area": "Build front",
        "purpose": "Trava versoes exatas das dependencias do projeto principal para instalacoes reproduziveis.",
    },
    {
        "path": "index.html",
        "area": "Front/PWA",
        "purpose": "HTML base servido pelo Vite. Declara root React, manifest, icones, viewport, preconnect Firebase e script src/main.jsx.",
    },
    {
        "path": "manifest.json",
        "area": "PWA legado",
        "purpose": "Copia do manifesto PWA na raiz. O Vite normalmente publica a versao de public/manifest.json.",
    },
    {
        "path": "vite.config.js",
        "area": "Build front",
        "purpose": "Configura Vite com React, Tailwind, alias @ para src e limite de aviso de chunk.",
    },
    {
        "path": "eslint.config.js",
        "area": "Qualidade",
        "purpose": "Configura ESLint flat config para JS/JSX, React Hooks, React Refresh e regra de variaveis nao usadas.",
    },
    {
        "path": "firebase.json",
        "area": "Deploy Firebase",
        "purpose": "Configura Hosting, headers de seguranca/cache, rewrites SPA, origem de Functions e caminhos de Firestore rules/indexes.",
    },
    {
        "path": "vercel.json",
        "area": "Deploy Vercel",
        "purpose": "Define rewrites para SPA e headers de seguranca/cache quando o app e publicado pela Vercel.",
    },
    {
        "path": "firestore.rules",
        "area": "Back/security",
        "purpose": "Regras de seguranca Firestore: users privados, candidatos publicos, choices publicas anonimizadas e colecoes administradas.",
    },
    {
        "path": "firestore.indexes.json",
        "area": "Back/data",
        "purpose": "Indices compostos para consultas/contagens em publicCandidateChoices por electionId, state e candidateIds.",
    },
]


LOG_FILES = [
    "codex-docs-vite.err.log",
    "codex-docs-vite.out.log",
    "codex-vite-5174.err.log",
    "codex-vite-5174.out.log",
    "codex-vite-cardstudy.err.log",
    "codex-vite-cardstudy.out.log",
    "codex-vite-share.err.log",
    "codex-vite-share.out.log",
    "codex-vite-verify.err.log",
    "codex-vite-verify.out.log",
    "codex-vite.err.log",
    "codex-vite.log",
    "vite-5174.err.log",
    "vite-5174.out.log",
    "vite-dev-5174.err.log",
    "vite-dev-5174.out.log",
]


SRC_CORE_FILES = [
    {
        "path": "src/main.jsx",
        "area": "Front entry",
        "purpose": "Ponto inicial do React. Importa CSS global, instala debug, registra service worker e renderiza UserProvider + App.",
    },
    {
        "path": "src/app/App.jsx",
        "area": "Front router",
        "purpose": "Organiza lazy loading das paginas, loading inicial, roteamento, redirects por login/desktop e montagem fixa do PrivacyConsent.",
    },
    {
        "path": "src/providers/UserProvider.jsx",
        "area": "Front state",
        "purpose": "Provider global de usuario: Firebase Auth, documento users/{uid}, elegibilidade, migracao de usuario e filtro persistido.",
    },
    {
        "path": "src/contexts/UserContext.js",
        "area": "Front state",
        "purpose": "Cria o contexto React usado para disponibilizar usuario, elegibilidade, loading e filtro ativo.",
    },
    {
        "path": "src/hooks/useUser.js",
        "area": "Front hook",
        "purpose": "Hook fino que encapsula useContext(UserContext).",
    },
    {
        "path": "src/hooks/useDesktopExperience.js",
        "area": "Front hook",
        "purpose": "Observa media query min-width: 768px para ativar experiencia publica desktop e reagir a resize.",
    },
]


PAGE_FILES = [
    {
        "path": "src/pages/Login.jsx",
        "area": "Tela",
        "purpose": "Tela de login mobile/entrada: Google popup, feedback por FlowToast, merge do rascunho visitante e redirect.",
    },
    {
        "path": "src/pages/Login.css",
        "area": "CSS tela",
        "purpose": "Layout do login, marca, video placeholder, CTA principal e ajustes mobile/desktop.",
    },
    {
        "path": "src/pages/Home.jsx",
        "area": "Tela",
        "purpose": "Tela de selecao de estado. Filtra UFs, salva estado local/remoto, confirma troca quando ja existem escolhas e abre tour.",
    },
    {
        "path": "src/pages/EscolherCandidatos.jsx",
        "area": "Tela",
        "purpose": "Tela de escolha por cargo. Busca candidatos, aplica filtros, restaura rascunho, calcula destaques/tallies e persiste selecao.",
    },
    {
        "path": "src/pages/MeuPlano.jsx",
        "area": "Tela",
        "purpose": "Tela de resumo do plano. Junta draft local/remoto, detalhes dos candidatos, media de nota/viabilidade, login/logout e compartilhamento.",
    },
    {
        "path": "src/pages/MeuPlano.css",
        "area": "CSS tela",
        "purpose": "Estilo da tela Meu Plano: header, perfil, cards de resumo, listas de escolhas, convite de login e responsividade.",
    },
    {
        "path": "src/pages/ContinuarPlano.jsx",
        "area": "Tela",
        "purpose": "Resgata token de QR Code/desktop para celular, exige login, salva draft na conta e direciona para proxima etapa.",
    },
    {
        "path": "src/pages/ContinuarPlano.css",
        "area": "CSS tela",
        "purpose": "Layout do painel de continuar plano, estados de loading/erro e botoes.",
    },
    {
        "path": "src/pages/LegalPage.jsx",
        "area": "Tela",
        "purpose": "Renderiza paginas legais e landing Sobre Nos. Usa conteudo de constants/legalPages e CTAs para comecar/login.",
    },
    {
        "path": "src/pages/LegalPage.css",
        "area": "CSS tela",
        "purpose": "Estilo da landing institucional e das paginas legais, incluindo hero, grids, secoes e breakpoints.",
    },
]


COMPONENT_FILES = [
    {
        "path": "src/components/selection/SelectBase.jsx",
        "area": "Componente UI",
        "purpose": "Shell reutilizavel de selecao: busca, lista/grid, cards, coluna de rascunho, filtros, continuar, modais e navegacao.",
    },
    {
        "path": "src/components/selection/SelectBase.css",
        "area": "CSS componente",
        "purpose": "Maior folha visual do fluxo de selecao: layout desktop/mobile, cards, header, filtros, rodape, sidebar e media queries.",
    },
    {
        "path": "src/components/selection/CandidateCard.jsx",
        "area": "Componente UI",
        "purpose": "Card de candidato com nome, partido, numero, nota, termometro de viabilidade, badges, modo compacto/detalhado e estado bloqueado.",
    },
    {
        "path": "src/components/navigation/BottomNavigation.jsx",
        "area": "Componente UI",
        "purpose": "Navegacao inferior/header das etapas. Le draft atual, calcula progresso e bloqueia etapas futuras quando faltam pre-requisitos.",
    },
    {
        "path": "src/components/navigation/BottomNavigation.css",
        "area": "CSS componente",
        "purpose": "Estilos da barra de etapas, icones, drawer/desktop variations, animacoes e breakpoints.",
    },
    {
        "path": "src/components/share/ShareChoicePanel.jsx",
        "area": "Componente UI",
        "purpose": "Painel/modal de compartilhamento: galeria arrastavel de templates, portal, acoes de compartilhar/salvar e status.",
    },
    {
        "path": "src/components/share/ShareChoicePanel.css",
        "area": "CSS componente",
        "purpose": "Estilos da galeria de artes, cards de preview, modal, acoes, dots e ajustes de altura/largura.",
    },
    {
        "path": "src/components/feedback/ConfirmModal.jsx",
        "area": "Componente UI",
        "purpose": "Modal generico de confirmacao/alerta com role dialog, titulo, mensagem, botoes e conteudo customizado.",
    },
    {
        "path": "src/components/feedback/ConfirmModal.css",
        "area": "CSS componente",
        "purpose": "Overlay, container, variacoes de tom, botoes e responsividade do ConfirmModal.",
    },
    {
        "path": "src/components/feedback/FlowToast.jsx",
        "area": "Componente UI",
        "purpose": "Toast simples com role=status para avisos de fluxo e mensagens temporarias.",
    },
    {
        "path": "src/components/feedback/FlowToast.css",
        "area": "CSS componente",
        "purpose": "Posicao, visual e keyframes de entrada/saida do FlowToast.",
    },
    {
        "path": "src/components/feedback/TourModal.jsx",
        "area": "Componente UI",
        "purpose": "Tour guiado: localiza elemento alvo, faz scrollIntoView, desenha destaque, tooltip e navegacao por passos.",
    },
    {
        "path": "src/components/feedback/TourModal.css",
        "area": "CSS componente",
        "purpose": "Overlay do tour, highlight box, tooltip, dots, animacoes fade e ajustes para telas pequenas.",
    },
    {
        "path": "src/components/privacy/PrivacyConsent.jsx",
        "area": "Componente UI",
        "purpose": "Banner de privacidade. Aceita cookies necessarios, expande detalhes e navega para personalizacao.",
    },
    {
        "path": "src/components/privacy/PrivacyConsent.css",
        "area": "CSS componente",
        "purpose": "Layout fixo do banner de privacidade, botoes e breakpoints mobile.",
    },
    {
        "path": "src/components/privacy/CookiePreferences.jsx",
        "area": "Componente UI",
        "purpose": "Tela/painel de switches de cookies: necessarios travados, opcionais editaveis, salvar/aceitar/recusar.",
    },
    {
        "path": "src/components/layout/AppFooter.jsx",
        "area": "Componente UI",
        "purpose": "Rodape institucional com links legais, contato e copyright.",
    },
    {
        "path": "src/components/layout/AppFooter.css",
        "area": "CSS componente",
        "purpose": "Estilo responsivo do rodape institucional e variacao para conteudo scrollavel.",
    },
    {
        "path": "src/components/icons/AppIcons.jsx",
        "area": "Componente UI",
        "purpose": "Biblioteca local de icones SVG: menu, info, clear, share, copy, download, search, filter, chevron, check, back e options.",
    },
    {
        "path": "src/components/icons/ChanceFlame.jsx",
        "area": "Componente UI",
        "purpose": "Componente de imagem para a chama da marca, reutilizado em loading, login, header e cards.",
    },
]


CONSTANT_FILES = [
    {
        "path": "src/constants/ballot.js",
        "area": "Constantes",
        "purpose": "ID da eleicao, nomes das Functions, versao do schema, rotas do fluxo, minimos por cargo e aliases legados.",
    },
    {
        "path": "src/constants/candidateRoutes.js",
        "area": "Constantes",
        "purpose": "Configura as rotas/titulos/chaves usadas para renderizar EscolherCandidatos para deputado e senadores.",
    },
    {
        "path": "src/constants/candidates.js",
        "area": "Constantes",
        "purpose": "Filtros da lista de candidatos e medias usadas na formula de viabilidade por cargo.",
    },
    {
        "path": "src/constants/states.js",
        "area": "Constantes",
        "purpose": "Lista oficial de UFs brasileiras usadas no Home, filtros e normalizacao.",
    },
    {
        "path": "src/constants/legalPages.js",
        "area": "Conteudo",
        "purpose": "Textos estruturados das paginas de cookies, privacidade, LGPD e informacoes institucionais.",
    },
]


SERVICE_FILES = [
    {
        "path": "src/services/firebase/firebase.js",
        "area": "Servico front",
        "purpose": "Inicializa Firebase client com VITE_*, fallback local, Firestore, Auth, Functions e GoogleAuthProvider.",
    },
    {
        "path": "src/services/voting/votingService.js",
        "area": "Servico front",
        "purpose": "Nucleo do rascunho/voto no cliente: localStorage, normalizacao de draft, publicCandidateChoices, handoff, recibos e callables.",
    },
    {
        "path": "src/services/candidates/candidateService.js",
        "area": "Servico front",
        "purpose": "Busca candidatos/partidos no Firestore, enriquece notas, cacheia dados e calcula contagens agregadas por candidato/estado.",
    },
    {
        "path": "src/services/share/shareCardService.js",
        "area": "Servico front",
        "purpose": "Gera analise, textos e imagens canvas 1080x1350 para compartilhar o plano; suporta Web Share, download e clipboard.",
    },
    {
        "path": "src/services/privacy/privacyPreferences.js",
        "area": "Servico front",
        "purpose": "Le e grava preferencias de privacidade/cookies em localStorage, incluindo migracao do consentimento legado.",
    },
    {
        "path": "src/services/pwa/registerServiceWorker.js",
        "area": "Servico PWA",
        "purpose": "Registra /sw.js apenas em producao e agenda update do service worker a cada 1 hora.",
    },
]


STYLE_UTIL_ASSET_FILES = [
    {
        "path": "src/styles/global.css",
        "area": "CSS global",
        "purpose": "Tokens e layout base da aplicacao: containers, loading, telas, utilitarios nv-*, cores e responsividade geral.",
    },
    {
        "path": "src/styles/reset.css",
        "area": "CSS global",
        "purpose": "Reset estrutural, box sizing, altura de app, foco visivel, touch behavior e prefers-reduced-motion.",
    },
    {
        "path": "src/styles/tailwind.css",
        "area": "CSS global",
        "purpose": "Entrada do Tailwind CSS v4 e tokens/compatibilidades usados pelo Vite.",
    },
    {
        "path": "src/utils/candidateMetrics.js",
        "area": "Utils front",
        "purpose": "Funcoes puras de nota, chance, tons de card e dados de exibicao de candidato.",
    },
    {
        "path": "src/utils/debugFlow.js",
        "area": "Utils front",
        "purpose": "Logs de fluxo com sanitizacao de dados sensiveis, historico local e ferramenta window.meuVotoDebug.",
    },
    {
        "path": "src/utils/search.js",
        "area": "Utils front",
        "purpose": "Normalizacao de busca e estado para comparacoes sem acento e com casing consistente.",
    },
    {
        "path": "src/utils/state.js",
        "area": "Utils front",
        "purpose": "Normaliza UFs/nomes de estados e extrai estado de candidatos a partir de campos variados.",
    },
    {
        "path": "src/assets/icone-fogo.png",
        "area": "Asset front",
        "purpose": "Icone raster de fogo usado como mascara/icone na navegacao e elementos visuais.",
    },
    {
        "path": "src/assets/nossovoto-192.png",
        "area": "Asset front",
        "purpose": "Icone oficial do nossovoto em 192x192 para reuso interno e documentacao visual.",
    },
    {
        "path": "src/assets/nossovoto-512.png",
        "area": "Asset front",
        "purpose": "Icone oficial do nossovoto em 512x512 para reuso interno e documentacao visual.",
    },
]


PUBLIC_FILES = [
    {
        "path": "public/manifest.json",
        "area": "PWA static",
        "purpose": "Manifesto servido em /manifest.json: nome, icones, modo standalone, orientacao e cores da PWA.",
    },
    {
        "path": "public/sw.js",
        "area": "PWA static",
        "purpose": "Service worker: caches de app/static, fallback offline para navegacao e limpeza de caches antigos.",
    },
    {
        "path": "public/icons/nossovoto-192.png",
        "area": "PWA asset",
        "purpose": "Icone 192x192 maskable/any para manifest, favicon e atalhos de instalacao.",
    },
    {
        "path": "public/icons/nossovoto-512.png",
        "area": "PWA asset",
        "purpose": "Icone 512x512 maskable/any para instalacao PWA e lojas/launchers.",
    },
]


FUNCTION_FILES = [
    {
        "path": "functions/package.json",
        "area": "Back build",
        "purpose": "Manifesto npm das Cloud Functions. Usa Node 20, firebase-admin, firebase-functions e scripts de deploy/lint.",
    },
    {
        "path": "functions/package-lock.json",
        "area": "Back build",
        "purpose": "Trava dependencias da pasta functions separadamente do front-end.",
    },
    {
        "path": "functions/index.js",
        "area": "Back serverless",
        "purpose": "Cloud Functions callable: valida payloads, salva draft remoto legado, gera/resgata handoff, apaga dados e registra voto anonimo.",
    },
]


DOC_FILES = [
    {
        "path": "docs/firestore-voting-model.md",
        "area": "Docs tecnica",
        "purpose": "Explica o modelo Firestore para escolhas publicas, choiceDocId, contagem e limites de privacidade.",
    },
    {
        "path": "docs/assets/nossovoto-512.png",
        "area": "Asset docs",
        "purpose": "Arquivo visual oficial do icone/logotipo guardado como referencia documental.",
    },
    {
        "path": "docs/generated-app-docs/build_documentation.py",
        "area": "Docs geradas",
        "purpose": "Script que gerou a documentacao tecnica visual anterior do app em DOCX/PDF.",
    },
    {
        "path": "docs/generated-app-docs/documentacao-app-nossovoto.docx",
        "area": "Docs geradas",
        "purpose": "Versao editavel da documentacao tecnica visual/front-end criada anteriormente.",
    },
    {
        "path": "docs/generated-app-docs/documentacao-app-nossovoto.pdf",
        "area": "Docs geradas",
        "purpose": "PDF da documentacao tecnica visual/front-end criada anteriormente.",
    },
    {
        "path": "docs/generated-app-docs/pdf-contact-sheet.png",
        "area": "QA docs",
        "purpose": "Folha de contato usada para revisar visualmente o PDF anterior.",
    },
    {
        "path": "docs/generated-app-docs/build_structure_documentation.py",
        "area": "Docs geradas",
        "purpose": "Este script: gera a documentacao de organizacao de pastas e arquivos.",
    },
    {
        "path": "docs/generated-app-docs/documentacao-estrutura-diretorios-nossovoto.docx",
        "area": "Docs geradas",
        "purpose": "Versao editavel desta documentacao de organizacao de pastas e arquivos.",
    },
    {
        "path": "docs/generated-app-docs/documentacao-estrutura-diretorios-nossovoto.pdf",
        "area": "Docs geradas",
        "purpose": "PDF final desta documentacao de organizacao de pastas e arquivos.",
    },
    {
        "path": "docs/generated-app-docs/structure-pdf-contact-sheet.png",
        "area": "QA docs",
        "purpose": "Folha de contato usada para revisar visualmente este PDF.",
    },
]


def add_title_page(doc: Document, package: dict) -> None:
    add_header_footer(doc)
    doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("nossovoto.org")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Documentacao de organizacao de pastas, diretorios e arquivos")
    add_callout(
        doc,
        "Escopo deste PDF",
        "Inventario tecnico do repositorio Plano-voto: o que e front-end, o que e back-end, como as pastas se conectam e o que cada arquivo de primeira mao faz. Pastas geradas ou de dependencia sao resumidas para manter a documentacao util.",
    )
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Projeto", package.get("name", "plano-mvp")],
            ["Versao", package.get("version", "-")],
            ["Data de geracao", datetime.now().strftime("%d/%m/%Y")],
            ["Raiz local", str(ROOT)],
            ["Front-end principal", "React 19 + Vite + React Router + Firebase client"],
            ["Back-end do projeto", "Firebase Functions + Firestore Rules/Indexes + Firebase managed services"],
        ],
        [2500, 6800],
    )
    doc.add_page_break()


def add_quick_map(doc: Document) -> None:
    doc.add_heading("1. Leitura rapida da organizacao", level=1)
    add_table(
        doc,
        ["Camada", "Onde fica", "Como pensar nela"],
        [
            ["Front-end / PWA", "src/, public/, index.html", "Tudo que roda no navegador: React, CSS, service worker, assets e camada cliente de Firebase."],
            ["Back-end serverless", "functions/index.js", "Codigo Node executado no Firebase Functions para operacoes sensiveis e transacionais."],
            ["Banco e seguranca", "firestore.rules, firestore.indexes.json", "Contrato de permissao e indices do Firestore; funciona como parte do back-end gerenciado."],
            ["Build/deploy", "package.json, vite.config.js, firebase.json, vercel.json", "Scripts, bundling, headers, rewrites e publicacao."],
            ["Documentacao", "README.md, CHANGELOG.md, docs/", "Explicacao humana e artefatos gerados para consulta."],
            ["Gerado/runtime", "dist/, node_modules/, functions/node_modules/, *.log", "Saida de build, dependencias e logs. Importante para rodar, mas nao e codigo-fonte de negocio."],
        ],
        [1900, 2500, 4900],
    )

    add_heading = doc.add_heading
    add_heading("1.1 Arvore principal", level=2)
    add_code_block(doc, build_tree_text())

    add_callout(
        doc,
        "Ponto importante",
        "A pasta src/services nao e back-end proprio: ela roda no front-end e apenas conversa com Firebase/Auth/Firestore/Functions, localStorage, canvas e service worker. O back-end escrito neste repositorio esta em functions/ e nas regras/indices Firestore.",
    )


def add_directory_inventory(doc: Document) -> None:
    doc.add_heading("2. Pastas do projeto", level=1)
    add_dir_table(
        doc,
        [
            ["src/", "Front-end", "Codigo principal da aplicacao React. Contem telas, componentes, contexto de usuario, hooks, servicos de cliente, CSS e utilitarios."],
            ["src/app/", "Front-end", "Camada de roteamento e inicializacao da experiencia dentro do React."],
            ["src/pages/", "Front-end", "Telas completas acessadas por rota: login, estado, candidatos, meu plano, continuar plano e paginas legais."],
            ["src/components/", "Front-end", "Componentes reutilizaveis, separados por dominio visual: feedback, selecao, navegacao, privacidade, compartilhamento, icones e layout."],
            ["src/services/", "Front-end/service layer", "Camada cliente que abstrai Firebase, cache local, candidatos, rascunho, compartilhamento, privacidade e PWA."],
            ["src/constants/", "Front-end/data", "Dados estaticos e configuracoes de fluxo usadas por varias telas."],
            ["src/styles/", "Front-end/CSS", "CSS global, reset e entrada Tailwind."],
            ["src/utils/", "Front-end/utils", "Funcoes puras compartilhadas por telas, services e componentes."],
            ["public/", "Front-end/PWA static", "Arquivos copiados diretamente para o build e servidos na raiz do site."],
            ["functions/", "Back-end serverless", "Cloud Functions Firebase em Node 20. Tem package separado e dependencias proprias."],
            ["docs/", "Documentacao", "Docs manuais, assets de referencia e PDFs/DOCX gerados."],
            ["dist/", "Gerado", "Build final do Vite. Deve ser recriado por npm run build, nao editado manualmente."],
            ["node_modules/", "Dependencias", "Dependencias npm do front-end. Deve ser recriado por npm install."],
            ["functions/node_modules/", "Dependencias back", "Dependencias npm das Cloud Functions. Deve ser recriado com npm install dentro de functions/."],
            [".git/", "Controle de versao", "Historico Git local. Nao e parte do app em producao."],
        ],
    )


def add_root_files(doc: Document) -> None:
    doc.add_heading("3. Arquivos da raiz", level=1)
    add_file_table(doc, ROOT_FILES)

    doc.add_heading("3.1 Logs e saidas locais", level=2)
    add_paragraph(
        doc,
        "Os arquivos abaixo sao saidas de execucoes locais do Vite/Codex. Eles ajudam a depurar uma sessao, mas nao fazem parte da arquitetura do app e podem ser apagados quando nao forem mais necessarios.",
    )
    add_file_table(
        doc,
        [{"path": path, "area": "Log local", "purpose": "Log de stdout/stderr de servidor Vite, teste visual ou execucao Codex local."} for path in LOG_FILES],
        [2800, 1200, 780, 720, 3800],
    )


def add_src_frontend(doc: Document) -> None:
    doc.add_heading("4. src/ - front-end React", level=1)
    add_callout(
        doc,
        "Fluxo mental do front-end",
        "src/main.jsx inicializa a aplicacao, UserProvider prepara usuario/elegibilidade, App.jsx decide as rotas, pages/ monta telas completas, components/ fornece blocos reutilizaveis e services/ centraliza acesso a Firebase/localStorage/canvas/PWA.",
    )

    doc.add_heading("4.1 Entrada, app, contexto e hooks", level=2)
    add_file_table(doc, SRC_CORE_FILES)

    doc.add_heading("4.2 pages/ - telas completas", level=2)
    add_file_table(doc, PAGE_FILES)

    doc.add_heading("4.3 components/ - objetos reutilizaveis", level=2)
    add_file_table(doc, COMPONENT_FILES)

    doc.add_heading("4.4 constants/ - dados e rotas estaticas", level=2)
    add_file_table(doc, CONSTANT_FILES)

    doc.add_heading("4.5 services/ - camada cliente", level=2)
    add_file_table(doc, SERVICE_FILES)

    doc.add_heading("4.6 styles/, utils/ e assets/", level=2)
    add_file_table(doc, STYLE_UTIL_ASSET_FILES)


def add_public_pwa(doc: Document) -> None:
    doc.add_heading("5. public/ - estaticos e PWA", level=1)
    add_file_table(doc, PUBLIC_FILES)
    add_bullets(
        doc,
        [
            "public/ e copiado pelo Vite para a raiz do build; por isso public/manifest.json vira /manifest.json em producao.",
            "public/sw.js roda no navegador como service worker. Ele nao e back-end, mas participa da disponibilidade offline e cache.",
            "public/icons/ contem os icones declarados no manifest e usados por instalacao PWA.",
        ],
    )


def add_backend(doc: Document) -> None:
    doc.add_heading("6. Back-end e dados", level=1)
    add_file_table(doc, FUNCTION_FILES)

    doc.add_heading("6.1 Funcoes exportadas em functions/index.js", level=2)
    add_table(
        doc,
        ["Function", "Papel"],
        [
            ["saveBallotState", "Salva ou reinicia o estado ativo do rascunho autenticado em transacao."],
            ["saveBallotStepSelection", "Valida candidatos por cargo/estado e grava grupo de candidatos no draft remoto legado."],
            ["deleteUserElectionData", "Remove draft/voto/elegibilidade legada e limpa campos do usuario."],
            ["createPlanHandoffToken", "Cria token temporario e hash para continuar plano desktop -> celular."],
            ["redeemPlanHandoffToken", "Valida token unico/expiracao e devolve draft para salvar na conta."],
            ["castAnonymousVote", "Valida elegibilidade, candidatos, estado e grava voto anonimo com recibo/auditoria/tallies."],
        ],
        [2600, 6700],
    )

    doc.add_heading("6.2 Regras e indices Firestore", level=2)
    add_table(
        doc,
        ["Arquivo", "Papel de back-end"],
        [
            ["firestore.rules", "Controla quem le/escreve users, choiceConfig, publicCandidateChoices, candidatos, partidos, elections, votes, drafts, tallies e handoff tokens."],
            ["firestore.indexes.json", "Da suporte a consultas agregadas count() em publicCandidateChoices usadas para viabilidade."],
        ],
        [2600, 6700],
    )


def add_docs_generated(doc: Document) -> None:
    doc.add_heading("7. docs/ - documentacao e artefatos", level=1)
    add_file_table(doc, DOC_FILES)
    add_dir_table(
        doc,
        [
            ["docs/assets/", "Assets docs", "Imagens de referencia usadas em documentacao ou identidade visual."],
            ["docs/generated-app-docs/", "Gerado", "Scripts, DOCX/PDF, screenshots e paginas rasterizadas usados para criar e validar documentacoes."],
            ["docs/generated-app-docs/screenshots/", "QA/gerado", "Capturas do app local usadas na documentacao tecnica visual anterior."],
            ["docs/generated-app-docs/pdf_pages/", "QA/gerado", "Paginas PNG geradas a partir do PDF para revisao visual."],
            ["docs/generated-app-docs/structure_pdf_pages/", "QA/gerado", "Paginas PNG geradas a partir deste PDF de estrutura para revisao visual."],
            ["docs/generated-app-docs/rendered/", "QA/gerado", "Saidas antigas do fluxo de render/validacao quando disponiveis."],
        ],
    )


def add_organization_logic(doc: Document) -> None:
    doc.add_heading("8. Como a organizacao se conecta", level=1)
    doc.add_heading("8.1 Cadeia de inicializacao", level=2)
    add_numbered(
        doc,
        [
            "index.html entrega um div#root e carrega /src/main.jsx.",
            "src/main.jsx importa CSS base, registra debug/PWA e monta UserProvider ao redor de App.",
            "UserProvider escuta Firebase Auth, carrega users/{uid}, eligibility e expõe dados por UserContext.",
            "App.jsx decide se mostra Login, Sobre Nos, Home, fluxo de candidatos, Meu Plano ou Continuar Plano.",
            "As pages chamam services/ para salvar estado, buscar candidatos, ler draft, gerar QR/share e tratar erros.",
            "components/ recebe props das pages e concentra comportamento visual reutilizavel.",
        ],
    )

    doc.add_heading("8.2 Fronteira front-end x back-end", level=2)
    add_table(
        doc,
        ["Item", "Classificacao correta", "Motivo"],
        [
            ["src/pages, src/components, src/styles", "Front-end", "Renderizam telas, estados visuais, CSS e interacoes no navegador."],
            ["src/services/firebase/candidates/voting/share", "Front-end service layer", "Codigo cliente. Chama Firebase/Firestore/Functions e usa browser APIs."],
            ["public/sw.js", "Front-end/PWA", "Roda no navegador como service worker e controla cache offline."],
            ["functions/index.js", "Back-end serverless", "Executa no Firebase Functions com firebase-admin e transacoes seguras."],
            ["firestore.rules/indexes", "Back-end gerenciado", "Define seguranca e estrutura de consulta no Firestore."],
            ["firebase.json/vercel.json/vite.config.js", "Infra/config", "Nao e regra de negocio direta; controla build, deploy, cache e headers."],
        ],
        [2400, 2100, 4800],
    )

    doc.add_heading("8.3 Onde editar cada tipo de coisa", level=2)
    add_table(
        doc,
        ["Quero mudar...", "Arquivo/pasta mais provavel"],
        [
            ["Nova rota ou redirect", "src/app/App.jsx e src/constants/ballot.js"],
            ["Tela de estado", "src/pages/Home.jsx e src/components/selection/SelectBase.jsx"],
            ["Tela de candidatos", "src/pages/EscolherCandidatos.jsx, SelectBase.jsx, CandidateCard.jsx e SelectBase.css"],
            ["Tela Meu Plano", "src/pages/MeuPlano.jsx, MeuPlano.css e ShareChoicePanel.*"],
            ["Cards de candidato", "src/components/selection/CandidateCard.jsx e SelectBase.css"],
            ["Calculo de nota/chance", "src/utils/candidateMetrics.js e src/services/candidates/candidateService.js"],
            ["Persistencia de rascunho/voto", "src/services/voting/votingService.js e functions/index.js"],
            ["Login/dados de usuario", "src/providers/UserProvider.jsx, src/pages/Login.jsx e firebase.js"],
            ["Privacidade/cookies", "privacyPreferences.js, PrivacyConsent.jsx, CookiePreferences.jsx e legalPages.js"],
            ["PWA/cache/offline", "public/sw.js, registerServiceWorker.js, public/manifest.json e index.html"],
            ["Regras de leitura/escrita", "firestore.rules e firestore.indexes.json"],
            ["Deploy", "firebase.json, vercel.json, package.json e functions/package.json"],
        ],
        [2900, 6400],
    )


def add_css_summary(doc: Document) -> None:
    doc.add_heading("9. Inventario CSS e responsividade por arquivo", level=1)
    css_files = [
        "src/styles/global.css",
        "src/styles/reset.css",
        "src/styles/tailwind.css",
        "src/pages/Login.css",
        "src/pages/ContinuarPlano.css",
        "src/pages/LegalPage.css",
        "src/pages/MeuPlano.css",
        "src/components/selection/SelectBase.css",
        "src/components/navigation/BottomNavigation.css",
        "src/components/share/ShareChoicePanel.css",
        "src/components/feedback/ConfirmModal.css",
        "src/components/feedback/FlowToast.css",
        "src/components/feedback/TourModal.css",
        "src/components/privacy/PrivacyConsent.css",
        "src/components/layout/AppFooter.css",
    ]
    rows = []
    for path in css_files:
        full_path = ROOT / path
        text = full_path.read_text(encoding="utf-8", errors="ignore") if full_path.exists() else ""
        media_count = text.count("@media")
        keyframe_count = text.count("@keyframes")
        rows.append([
            rel(path),
            text_file_line_count(full_path),
            str(media_count),
            str(keyframe_count),
            {
                "src/styles/global.css": "Base visual, loading, containers, utilitarios e tokens globais.",
                "src/styles/reset.css": "Reset, acessibilidade de foco, touch e reduced motion.",
                "src/styles/tailwind.css": "Entrada Tailwind e tokens iniciais.",
                "src/pages/Login.css": "Login e CTA inicial.",
                "src/pages/ContinuarPlano.css": "Painel de handoff QR.",
                "src/pages/LegalPage.css": "Landing Sobre Nos e paginas legais.",
                "src/pages/MeuPlano.css": "Resumo do plano e perfil.",
                "src/components/selection/SelectBase.css": "Layout principal de selecao, responsividade mais complexa e cards.",
                "src/components/navigation/BottomNavigation.css": "Barra de etapas, drawer e estados ativo/completo/futuro.",
                "src/components/share/ShareChoicePanel.css": "Modal/galeria de compartilhamento.",
                "src/components/feedback/ConfirmModal.css": "Modais de confirmacao.",
                "src/components/feedback/FlowToast.css": "Toast e animacoes.",
                "src/components/feedback/TourModal.css": "Tour guiado e highlight.",
                "src/components/privacy/PrivacyConsent.css": "Banner de privacidade.",
                "src/components/layout/AppFooter.css": "Rodape institucional.",
            }[path],
        ])
    add_table(doc, ["CSS", "Linhas", "@media", "@keyframes", "Responsabilidade"], rows, [2950, 700, 720, 900, 4030])


def add_generated_runtime(doc: Document) -> None:
    doc.add_heading("10. Pastas geradas, dependencias e cuidados", level=1)
    add_dir_table(
        doc,
        [
            ["dist/", "Build gerado", "Resultado de npm run build. Publicado em hosting, mas nao deve ser editado manualmente."],
            ["node_modules/", "Dependencias front", "Pacotes npm instalados para o app principal. Pode ser recriado com npm install."],
            ["functions/node_modules/", "Dependencias back", "Pacotes npm usados pelas Cloud Functions. Pode ser recriado dentro de functions/."],
            ["*.log", "Logs locais", "Saidas de servidores/execucoes locais. Servem para debug, nao para codigo-fonte."],
            ["docs/generated-app-docs/pdf_pages/", "QA gerado", "Imagens rasterizadas para revisar PDFs. Podem ser regeneradas."],
            ["docs/generated-app-docs/structure_pdf_pages/", "QA gerado", "Imagens rasterizadas desta documentacao de estrutura. Podem ser regeneradas."],
            ["docs/generated-app-docs/screenshots/", "QA gerado", "Capturas do app usadas como evidencia visual nas documentacoes."],
        ],
    )

    add_callout(
        doc,
        "Cuidados praticos",
        "Nao coloque valores de .env.local em docs ou commits. Para alterar negocio, edite src/ e functions/; para alterar build/deploy, edite os arquivos de configuracao da raiz; para alterar visual, procure primeiro CSS do componente ou da page responsavel.",
    )


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    package = json.loads((ROOT / "package.json").read_text(encoding="utf-8"))
    doc = Document()
    style_document(doc)
    add_title_page(doc, package)
    add_quick_map(doc)
    add_directory_inventory(doc)
    add_root_files(doc)
    add_src_frontend(doc)
    add_public_pwa(doc)
    add_backend(doc)
    add_docs_generated(doc)
    add_organization_logic(doc)
    add_css_summary(doc)
    add_generated_runtime(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
