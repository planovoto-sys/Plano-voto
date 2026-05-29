from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "generated-app-docs"
SCREENSHOT_DIR = OUT_DIR / "screenshots"
DOCX_PATH = OUT_DIR / "documentacao-app-nossovoto.docx"

APP_ORANGE = RGBColor(255, 152, 0)
APP_ORANGE_DARK = RGBColor(189, 100, 0)
INK = RGBColor(17, 17, 17)
MUTED = RGBColor(87, 87, 94)
BLUE = RGBColor(46, 116, 181)
BLUE_DARK = RGBColor(31, 77, 120)
LIGHT_GRAY = "F2F4F7"
SOFT_ORANGE = "FFF4E0"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DADCE0", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, color=None, font="Calibri", italic=None) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE_DARK, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = styles.add_style("CodeBlock", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(40, 44, 52)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "nossovoto.org | Documentacao tecnica do app"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Gerado em 28/05/2026 | Fonte: repositorio local Plano-voto"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.runs[0], size=9, color=MUTED)


def add_title_page(doc: Document, metadata: dict) -> None:
    for _ in range(3):
        doc.add_paragraph()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("nossovoto.org")
    set_run_font(run, size=18, bold=True, color=APP_ORANGE_DARK, font="Calibri")

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Documentacao tecnica e visual do app")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Logicas de front-end, responsividade, transicoes, efeitos, telas, objetos, dados e fluxo de navegacao"
    )
    set_run_font(subtitle_run, size=13, color=MUTED)

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    rows = [
        ("Projeto", "Plano-voto / nossovoto.org"),
        ("Versao do package.json", metadata.get("version", "1.5.1")),
        ("Stack principal", "React 19, Vite 7, Firebase, React Router, CSS modular por tela/componente"),
        ("Data da documentacao", "28/05/2026"),
        ("Verificacao executada", "npm run build concluido com sucesso"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_width(row.cells[0], 2500)
        set_cell_width(row.cells[1], 6860)
        for cell in row.cells:
            set_cell_margins(cell)
            set_cell_border(cell, "DADCE0", "6")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                set_run_font(r, size=10.5, bold=True, color=INK)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                set_run_font(r, size=10.5, color=INK)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "Escopo: esta documentacao descreve o estado atual do codigo local e nao substitui revisao juridica, eleitoral ou de seguranca."
    )
    set_run_font(note_run, size=10, italic=True, color=MUTED)
    doc.add_page_break()


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_code(doc: Document, code: str) -> None:
    for line in code.strip().splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        set_cell_border(hdr[idx])
        set_cell_margins(hdr[idx])
        if widths:
            set_cell_width(hdr[idx], widths[idx])
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                set_run_font(r, size=9.5, bold=True, color=INK)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            if widths:
                set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_font(r, size=9.3, color=INK)


def add_screenshot(doc: Document, filename: str, caption: str, width=6.1) -> None:
    path = SCREENSHOT_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, italic=True, color=MUTED)


def add_section_intro(doc: Document) -> None:
    doc.add_heading("1. Visao geral", level=1)
    add_paragraph(
        doc,
        "O app e um PWA de planejamento de voto. O usuario escolhe estado, monta um rascunho com deputado federal e senadores, revisa o plano e pode salvar/compartilhar quando estiver autenticado. O codigo esta organizado para ter uma experiencia publica de exploracao em telas desktop e uma entrada mobile centrada em login/continuidade."
    )
    add_bullets(
        doc,
        [
            "Produto: nossovoto.org, um fluxo guiado para organizar escolhas de voto por UF e cargo.",
            "Front-end: React 19 com rotas lazy-loaded, React Router 7, hooks proprios e CSS por tela/componente.",
            "Dados: Firebase Auth, Cloud Firestore, Callable Cloud Functions e cache local via localStorage.",
            "PWA: manifest, service worker em producao, cache de shell, assets e fallback de navegacao.",
            "Design: mobile-first ate 767px; layout amplo a partir de 768px; comportamento de escritorio completo a partir de 1024px.",
        ],
    )

    add_table(
        doc,
        ["Camada", "Arquivos principais", "Responsabilidade"],
        [
            ["Bootstrap", "src/main.jsx, src/app/App.jsx", "Carregar estilos globais, instalar debug, registrar PWA, montar provider e rotas."],
            ["Estado do usuario", "UserProvider.jsx, UserContext.js, useUser.js", "Sincronizar Firebase Auth, perfil Firestore, elegibilidade e filtro ativo."],
            ["Fluxo de voto", "Home.jsx, EscolherCandidatos.jsx, MeuPlano.jsx", "Estado, selecao por cargo, rascunho final, compartilhamento e login quando necessario."],
            ["Persistencia", "votingService.js, candidateService.js", "Normalizar rascunhos, salvar local/remoto, buscar candidatos, cachear listas e recalcular viabilidade."],
            ["Interface", "SelectBase.jsx, CandidateCard.jsx, BottomNavigation.jsx, ShareChoicePanel.jsx", "Objetos visuais reutilizaveis, navegacao, cards, filtros, modais e galeria de compartilhamento."],
            ["Infra", "functions/index.js, firestore.rules, public/sw.js", "Funcoes callable, regras de seguranca, tokens de handoff e cache offline."],
        ],
        [1500, 2500, 5360],
    )

    doc.add_heading("2. Estrutura de pastas", level=1)
    add_table(
        doc,
        ["Pasta/arquivo", "Papel no app"],
        [
            ["src/app", "Camada de roteamento e decisoes iniciais de entrada."],
            ["src/pages", "Telas de login, landing/legal, selecao de estado, candidatos, meu plano e handoff."],
            ["src/components/selection", "Base visual do fluxo de selecao e cards de candidatos."],
            ["src/components/feedback", "Modal de confirmacao, toast de fluxo e tour guiado."],
            ["src/components/share", "Painel e modal de compartilhamento, com templates de arte."],
            ["src/components/navigation", "Navegacao inferior/mobile e navegacao de etapas no header/desktop."],
            ["src/services", "Firebase, voto, candidatos, privacidade, PWA e arte de compartilhamento."],
            ["functions", "Cloud Functions callable para voto, exclusao, token de continuidade e resgate."],
            ["public", "Manifest, service worker e icones da PWA."],
            ["docs", "Documentos tecnicos e artefatos gerados."],
        ],
        [2500, 6860],
    )


def add_routing_and_boot(doc: Document) -> None:
    doc.add_heading("3. Bootstrap, loading e roteamento", level=1)
    add_paragraph(
        doc,
        "O ponto de entrada importa Tailwind, reset e CSS global, instala ferramentas de debug do fluxo, registra o service worker em producao e renderiza App dentro do UserProvider. O App usa lazy imports para as paginas e Suspense com uma tela de carregamento visual."
    )
    add_code(
        doc,
        """
main.jsx -> installFlowDebugTools()
main.jsx -> registerPwaServiceWorker()
main.jsx -> <UserProvider><App /></UserProvider>
App.jsx  -> BrowserRouter + Routes + Suspense + PrivacyConsent
        """,
    )
    add_table(
        doc,
        ["Decisao", "Como funciona"],
        [
            ["Intro visual", "introReady so libera as rotas depois de 1500ms, evitando troca brusca entre estado de auth e primeira tela."],
            ["Preload de rotas", "Depois que auth termina, requestIdleCallback pre-carrega Login ou Home/EscolherCandidatos/MeuPlano."],
            ["Desktop publico", "useDesktopExperience considera min-width: 768px. Em desktop, a raiz mostra a landing Sobre e libera exploracao publica."],
            ["Mobile sem login", "A raiz mostra Login. Rotas publicExplorationRoute redirecionam para / quando nao ha user e nao e desktop."],
            ["Retomada autenticada", "AuthenticatedEntryRedirect le rascunho local/remoto, calcula progresso e navega para a proxima etapa."],
            ["Consentimento", "PrivacyConsent e montado fora de Routes para aparecer sobre qualquer tela ate o aceite."],
        ],
        [2200, 7160],
    )

    doc.add_heading("4. Mapa de rotas", level=1)
    add_table(
        doc,
        ["Rota", "Tela/componente", "Regra de acesso/uso"],
        [
            ["/", "Login ou Sobre ou redirect de rascunho", "Mobile visitante: Login. Desktop: Sobre. Usuario logado: retoma proxima etapa."],
            ["/login", "Login", "Entrada Google, merge de rascunho visitante e retorno para rota original."],
            ["/home", "Home + SelectBase home-state", "Escolha de estado. Publica no desktop, autenticada no mobile."],
            ["/escolher-deputado-federal", "EscolherCandidatos", "Lista Deputado Federal e salva grupo deputado_federal."],
            ["/escolher-senadores", "EscolherCandidatos", "Lista Senadores unificada, minimo 2, salva senadores_1 e zera senadores_2 legado."],
            ["/meu-plano", "MeuPlano", "Resumo do rascunho, indicadores medios e compartilhamento quando completo/logado."],
            ["/continuar-plano/:token", "ContinuarPlano", "Resgata QR/token temporario e salva draft na conta autenticada."],
            ["/cookies", "LegalPage cookies + CookiePreferences", "Gerenciamento de preferencias de privacidade."],
            ["/politica-de-privacidade", "LegalPage privacidade", "Conteudo institucional/legal."],
            ["/lgpd", "LegalPage LGPD", "Conteudo institucional/legal."],
            ["/sobre-nos", "LegalPage type=sobre", "Landing de produto e explicacao do fluxo."],
            ["Rotas legadas", "/meu-voto, /meuvoto, /resultado, /finalizacao", "Redirecionam para /meu-plano via privateRedirect."],
        ],
        [2300, 2600, 4460],
    )


def add_data_logic(doc: Document) -> None:
    doc.add_heading("5. Modelo de dados e persistencia", level=1)
    add_paragraph(
        doc,
        "A entidade central do front-end e o BallotDraft: um rascunho normalizado que existe no localStorage para navegacao rapida e, quando o usuario faz login, e sincronizado com um documento publico aleatorio no Firestore."
    )
    add_code(
        doc,
        """
BallotDraft {
  schema_version: 1,
  election_id: "congresso-2026",
  estado: "AC" | "SP" | null,
  selections: { deputado_federal: [], senadores: [] },
  candidate_groups: { deputado_federal: [], senadores_1: [], senadores_2: [] },
  completed_steps: { deputado_federal: false, senadores_1: false, senadores_2: false },
  updated_at: ISODate | null
}
        """,
    )
    add_table(
        doc,
        ["Objeto/colecao", "Detalhe tecnico"],
        [
            ["localStorage", "Prefixo meuvoto:congresso-2026. Guarda ballotDraft por usuario e visitor:local. Rascunhos expiram localmente apos 30 dias."],
            ["users/{uid}", "Perfil minimo: nome, email, imagem, estado, role voter, schema_version e timestamps. Nao guarda a lista publica de candidatos."],
            ["users/{uid}/private/choiceConfig", "Documento privado com choiceDocId aleatorio, usado para desvincular o UID do documento publico."],
            ["publicCandidateChoices/{choiceDocId}", "Documento publico de contagem com electionId, state, candidateIds e updatedAt. Lido por count() para viabilidade."],
            ["candidatos", "Fonte publica dos candidatos. O app tolera chaves legadas como Nome/Cargo/Partido e novas como nome/cargo/partido."],
            ["partidos_politicos", "Usado para enriquecer candidatos ingressantes com nota partidaria quando a nota do candidato nao existe."],
            ["plan_handoff_tokens", "Criado por Cloud Function; token e hasheado, expira em 10 minutos e e de uso unico."],
            ["votes/eligibility", "Fluxo de voto confirmado passa por Cloud Function; elegibilidade impede voto duplicado e retorna recibo."],
        ],
        [2700, 6660],
    )

    doc.add_heading("5.1 Viabilidade e notas", level=2)
    add_bullets(
        doc,
        [
            "calculateCandidateChance limita o resultado entre 0 e 100 e usa round((selecoesAtivas / mediaVotosEleitos) * 100).",
            "AVERAGE_ELECTED_VOTES_BY_OFFICE atualmente define 3 para deputado_federal e 3 para senadores.",
            "getCandidateSystemScore prioriza nota_final/notaFinal/notaCandidato/Nota candidato e, se necessario, nota partidaria.",
            "getCandidateTone transforma dados em estados visuais: success, danger, new, neutral ou visitor.",
            "Cards bloqueados em modo visitante escondem indicadores personalizados e convidam o usuario a fazer login.",
        ],
    )

    doc.add_heading("5.2 Cache de candidatos e contagens", level=2)
    add_table(
        doc,
        ["Cache", "Janela fresca", "Janela stale", "Uso"],
        [
            ["Candidatos por cargo", "12h", "7 dias", "Renderiza rapido com cache local e depois atualiza por rede quando necessario."],
            ["Partidos", "12h", "7 dias", "Enriquece ingressantes com nota partidaria."],
            ["Tallies por candidato/estado", "2 min", "30 min", "Evita count() repetido em cada interacao."],
            ["Contagens por estado", "2 min", "30 min", "Base pronta para mostrar atividade agregada por UF."],
        ],
        [2300, 1700, 1700, 3660],
    )

    doc.add_heading("5.3 Regras de seguranca", level=2)
    add_bullets(
        doc,
        [
            "Somente o dono le users/{uid} e users/{uid}/private/choiceConfig; admin pode listar/alterar.",
            "Candidatos e partidos sao publicos para leitura; escrita e restrita a admin.",
            "publicCandidateChoices e publico para leitura, mas criacao/atualizacao exige que o choiceDocId pertenca ao usuario autenticado.",
            "votes, audit_events e plan_handoff_tokens nao podem ser escritos diretamente pelo cliente.",
            "Headers de deploy bloqueiam iframe via X-Frame-Options DENY, usam nosniff, Referrer-Policy restrita e Permissions-Policy sem camera/microfone/geolocalizacao.",
        ],
    )


def add_screen_docs(doc: Document) -> None:
    doc.add_heading("6. Telas e objetos visuais", level=1)
    add_paragraph(
        doc,
        "As telas foram documentadas observando codigo e capturas reais do app local. Os objetos abaixo sao os blocos que o usuario percebe e que o front-end controla como componentes."
    )

    doc.add_heading("6.1 Landing / Sobre nos", level=2)
    add_screenshot(doc, "desktop-sobre-nos.png", "Tela Sobre em desktop: primeira experiencia publica e explicacao do produto.")
    add_table(
        doc,
        ["Objeto", "Logica/funcao"],
        [
            ["Header institucional", "Marca com ChanceFlame, links de ancora e CTA de login. Em mobile inclui acao Voltar quando aplicavel."],
            ["Hero", "Texto de valor, CTA comecar sem login e CTA login; no desktop o preview visual simula estado, deputado, senadores e rascunho."],
            ["Preview do produto", "Cards estaticos usados como sinal visual do fluxo sem depender de dados reais."],
            ["Secoes Como funciona", "Trio de passos: estado, rascunho, salvar/continuar."],
            ["Comparativo sem/com login", "Mostra diferenca de permissao e valor entre visitante e conta autenticada."],
            ["Privacidade/celular", "Explica rascunho local, conta e QR Code temporario."],
        ],
        [2300, 7060],
    )

    doc.add_heading("6.2 Login", level=2)
    add_screenshot(doc, "mobile-login.png", "Login em viewport mobile: marca, card de video futuro e CTA principal.", width=2.65)
    add_table(
        doc,
        ["Objeto", "Estados e efeitos"],
        [
            ["Marca nossovoto.org", "H1 acessivel com aria-label e icone ChanceFlame."],
            ["Video-card", "Placeholder com botao play; hoje dispara FlowToast 'Video em breve.'."],
            ["Botao COMEÇAR", "Chama signInWithPopup, bloqueia reentrancia via loginSubmitting e muda texto para ENTRANDO."],
            ["Merge de rascunho", "Apos login, mergeVisitorBallotDraftIntoAccount salva o draft visitante na conta e limpa visitor:local."],
            ["Erros", "FlowToast informa login indisponivel, falha Google ou merge local nao salvo."],
            ["Retorno", "location.state.from define a rota de volta; fallback e /."],
        ],
        [2300, 7060],
    )

    doc.add_heading("6.3 Selecao de estado", level=2)
    add_screenshot(doc, "desktop-home-estado.png", "Selecao de estado em desktop: busca, grid e navegacao de etapas no header.")
    add_table(
        doc,
        ["Objeto", "Logica/funcao"],
        [
            ["SelectBase variant home-state", "Reusa o motor de selecao com limite 1, autoAvancarAoSelecionar e copy reduzida para 'Estado'."],
            ["Busca", "useDeferredValue + normalizeSearch filtram nome, sigla e nome completo."],
            ["Grid de estados", "27 UFs de BRAZILIAN_STATES; cada card e um botao com nome e sigla."],
            ["Troca de estado", "Se ja existe draft com escolhas e o usuario troca UF, ConfirmModal alerta que selecoes atuais serao apagadas."],
            ["Persistencia visitante", "saveVisitorBallotState grava localStorage. Usuario logado usa saveBallotState e documento publico."],
            ["Navegacao", "Ao confirmar, navega para BALLOT_ROUTES.deputadoFederal com bypassVoteRedirect."],
        ],
        [2300, 7060],
    )

    doc.add_heading("6.4 Selecao de candidatos", level=2)
    add_screenshot(doc, "desktop-escolher-deputado.png", "Lista de deputado federal: coluna de rascunho, busca, filtro e cards com termometro.")
    add_table(
        doc,
        ["Objeto", "Logica/funcao"],
        [
            ["Candidate flow", "Duas areas no desktop: 'Meu Candidato' e 'Candidatos'. Em mobile vira fluxo vertical."],
            ["Busca de candidatos", "Filtra por nome, partido e numero usando busca diferida para reduzir renders durante digitacao."],
            ["Filtro", "Dropdown com Todos, Atuais, Novos e Selecionados; fecha em clique externo ou Escape."],
            ["Cards", "CandidateCard recebe tone, selected, featuredMetrics, displayMode e estado bloqueado visitante."],
            ["Termometro", "Barra horizontal com progresso por --thermometer-progress, ticks e label como META ATINGIDA, INTERESSANTE ou NOTA BAIXA."],
            ["Destaque", "featuredCandidateId prioriza score > 7, chance > 0 e < 100; desempata por chance, nota e nome."],
            ["Modais de decisao", "Nota baixa pede MUDAR/MANTER; viabilidade 100 pede manter/trocar; duplicado impede repetir candidato em outra etapa."],
            ["Persistencia otimista", "handleSelectionChange atualiza UI, salva draft, aplica delta local de tallies e depois tenta refresh do servidor."],
            ["Render incremental", "Mostra os primeiros 80 candidatos e oferece 'Mostrar todos' para evitar carga inicial pesada."],
        ],
        [2300, 7060],
    )

    doc.add_heading("6.5 Meu Plano", level=2)
    add_screenshot(doc, "desktop-meu-plano.png", "Meu Plano em desktop: resumo, candidatos selecionados e convite de login/compartilhamento.")
    add_table(
        doc,
        ["Objeto", "Logica/funcao"],
        [
            ["Header", "Avatar/perfil, marca central e botao de login/logout. O popover mostra nome e email/plano local."],
            ["Overview tiles", "Estado, media de nota e media de viabilidade calculadas a partir dos candidatos escolhidos."],
            ["Choice sections", "Deputado Federal e Senadores usam CandidateCard em modo summary e levam de volta para edicao."],
            ["EmptyChoiceCard", "Quando falta cargo, mostra CTA 'Escolher' e explicacao curta."],
            ["Guest mode", "Sem user.uid, mostra 'Entrar para compartilhar' e informa que o plano esta salvo no dispositivo."],
            ["ShareChoicePanel", "So e liberado para usuario logado com deputado + 2 senadores + estado."],
            ["Detalhamento remoto", "MeuPlano busca draft remoto e detalhes atualizados de candidatos/tallies; cai para snapshot local se falhar."],
        ],
        [2300, 7060],
    )

    doc.add_heading("6.6 Continuar plano via QR", level=2)
    add_table(
        doc,
        ["Estado", "Comportamento"],
        [
            ["Sem login", "Mostra painel 'Continuar plano' e CTA Fazer login; o QR nao autentica automaticamente."],
            ["Logado/carregando", "Exibe loading com ChanceFlame enquanto redeemPlanHandoffToken e saveBallotDraftToAccount rodam."],
            ["Sucesso", "Calcula getBallotProgress(savedDraft) e navega para a proxima etapa com FlowToast 'Plano carregado neste celular.'."],
            ["Erro", "Informa que o QR expirou ou ja foi usado e mostra botao Voltar."],
            ["Seguranca", "Token gerado com 32 bytes, armazenado como hash, TTL de 10 minutos e marcado como usado no resgate."],
        ],
        [2200, 7160],
    )


def add_components(doc: Document) -> None:
    doc.add_heading("7. Componentes reutilizaveis", level=1)
    add_table(
        doc,
        ["Componente", "Responsabilidade", "Detalhes de estado/interacao"],
        [
            ["SelectBase", "Shell das telas de selecao", "Controla selecionados, filtros, busca, modais, botao continuar, modo compacto/detalhado e layout desktop/mobile."],
            ["CandidateCard", "Card de candidato", "Calcula score, partido, numero, status visual, termometro, bloqueio visitante e acao selecionar/remover."],
            ["BottomNavigation", "Etapas do voto", "Mostra Estado, Deputados, Senadores e NossoVoto; desabilita etapas futuras conforme progresso do draft."],
            ["ConfirmModal", "Dialogos de confirmacao", "role=dialog, aria-modal, titulo por useId, variantes perigo/low-score/high-chance/login-required."],
            ["TourModal", "Guia contextual", "Localiza target por querySelector, scrollIntoView, highlight animado e tooltip top/bottom."],
            ["FlowToast", "Mensagem temporaria", "role=status com aria-live polite; animacao de entrada e saida."],
            ["ShareChoicePanel", "Compartilhamento", "Modal com galeria de templates, download de imagem canvas, copiar texto e Web Share API."],
            ["PrivacyConsent", "Consentimento global", "Banner fixo com detalhes expansivel, personalizacao em /cookies e aceite de somente necessarios."],
            ["CookiePreferences", "Preferencias", "Switches para analytics, personalizacao, marketing e uso comercial agregado; necessarios sempre ativos."],
            ["AppFooter", "Rodape institucional", "Links legais, contato e copyright dentro do scroll principal."],
        ],
        [1900, 3000, 4460],
    )

    doc.add_heading("7.1 SelectBase por dentro", level=2)
    add_bullets(
        doc,
        [
            "Usa DESKTOP_LAYOUT_QUERY em 1024px para decidir se mostra sidebar de rascunho, botao continuar desktop e grids mais amplos.",
            "Sincroniza selecao inicial por assinatura de IDs + chance + score + destaque, evitando re-render inutil quando os mesmos candidatos permanecem.",
            "Mostra/oculta o botao Continuar conforme selecao minima e direcao do scroll: subir revela, descer esconde.",
            "Para senadores, permite selecionar mais de um, exige minimo 2 para navegar e mostra modal de substituicao quando limite explicito e atingido.",
            "A UI de filtro fecha com pointerdown externo e Escape, melhorando comportamento de dropdown em desktop.",
        ],
    )

    doc.add_heading("7.2 CandidateCard por dentro", level=2)
    add_bullets(
        doc,
        [
            "getSingleLineSize injeta variaveis CSS para reduzir textos longos como assessment sem quebrar o card.",
            "A classe final combina tone, selected, summary, fire-featured, viability-complete, blocked, expandable e compact/detailed.",
            "O termometro usa CSS custom properties para representar a porcentagem e evita recalcular geometria no DOM.",
            "Em modo expandivel, o clique abre painel de detalhes; em modo selecao, o mesmo botao chama onSelect.",
            "Campos personalizados bloqueados em visitante mostram insight de login em vez de nota/viabilidade real.",
        ],
    )

    doc.add_heading("7.3 Compartilhamento", level=2)
    add_bullets(
        doc,
        [
            "Templates: resumo, completo, termometro e checklist. Cada um tem label, thumbnail, tag e descricao.",
            "createShareAnalysis calcula media de nota, media de chance, perfil do plano e indicadores para a arte.",
            "A imagem final e desenhada em canvas 1080x1350, com fundo, paineis, pills, barras de progresso e linhas de texto quebradas por largura.",
            "shareTemplate tenta Web Share API com arquivo PNG; se o navegador nao aceitar arquivos, compartilha texto/url; se nao houver Web Share, copia texto.",
        ],
    )


def add_responsiveness(doc: Document) -> None:
    doc.add_heading("8. Responsividade e layout", level=1)
    add_paragraph(
        doc,
        "A base visual e mobile-first. O CSS global fixa o app como uma superficie controlada por --app-width e --app-height: 100dvh, enquanto breakpoints ampliam para tela cheia e reorganizam header, grids e navegacao."
    )
    add_table(
        doc,
        ["Breakpoint", "Mudanca principal"],
        [
            ["<= 320px", "Reduz gutter para 10px, fontes de header e dimensoes de CandidateCard para caber em aparelhos estreitos."],
            ["<= 360px", "Gutter 14px, botoes de header 40px, cards e badges menores."],
            ["<= 420px", "Gutter 18px e ajustes de tipografia; Login e Share modal reduzem paddings."],
            ["431px-767px", "App deixa de ficar preso aos 430px e ocupa 100% da largura; ainda usa navegacao de rodape."],
            [">= 768px", "Desktop/tablet: #root ocupa 100%, prototype-page vira grid header/main/footer, scrollbars finas aparecem, grids usam 2 colunas ou auto-fill."],
            [">= 1024px", "Experiencia desktop completa: bottom nav footer some, navegacao de etapas vai ao header, sidebar de rascunho fica visivel e sticky em areas apropriadas."],
            [">= 1280px", "Candidatos e estados ganham colunas maiores, gutters aumentam e header distribui marca/etapas/acoes com mais respiro."],
            [">= 1536px / 1920px", "Gutters sobem por clamp para evitar conteudo estourado em monitores largos."],
        ],
        [2200, 7160],
    )

    doc.add_heading("8.1 Padroes responsivos recorrentes", level=2)
    add_bullets(
        doc,
        [
            "safe-area: rodape, modal de compartilhamento, tour e consentimento usam env(safe-area-inset-bottom/top).",
            "Min-width defensivo: quase todos os containers importantes usam min-width: 0 para evitar overflow em grids/flex.",
            "Grid adaptativo: cards usam repeat(auto-fill/auto-fit, minmax(min(100%, Xpx), 1fr)).",
            "Header: mobile centraliza titulo e mostra icones; desktop adiciona marca a esquerda, progresso a direita e acoes no extremo.",
            "Scroll: areas principais escondem scrollbar no mobile e mostram scrollbar fina no desktop.",
            "Tipografia: o projeto evita letter-spacing negativo e usa variaveis/clamp em cards desktop para estabilidade visual.",
        ],
    )

    doc.add_heading("8.2 Navegacao", level=2)
    add_table(
        doc,
        ["Contexto", "Comportamento"],
        [
            ["Mobile/tablet ate 1023px", "BottomNavigation e footer fixo: 4 itens com icones, altura 58px + safe-area."],
            ["Desktop >=1024px", "BottomNavigation placement header fica visivel; footer placement e escondido por --bottom-nav-height: 0."],
            ["Etapas futuras", "Deputado exige estado; Senadores exige deputado; Meu Plano continua acessivel para revisao."],
            ["Indicador ativo", "Classe is-active e pseudo-elemento/linha laranja marcam a etapa atual."],
            ["Rotas legadas", "A navegacao considera aliases para meu voto/resultado/finalizacao como NossoVoto."],
        ],
        [2300, 7060],
    )


def add_transitions(doc: Document) -> None:
    doc.add_heading("9. Transicoes, efeitos e estados visuais", level=1)
    add_table(
        doc,
        ["Efeito", "Onde aparece", "Implementacao"],
        [
            ["Loading intro", "App, SelectBase, MeuPlano, ContinuarPlano", "ChanceFlame central com drop-shadow; CSS possui keyframes loadingSpin, loadingPulse, loadingSweep e loadingInnerSweep, embora partes de ring/scan estejam desativadas."],
            ["Reduced motion", "reset.css e global.css", "prefers-reduced-motion zera duracao de animacoes/transicoes e remove animacoes do loading."],
            ["Toast", "FlowToast", "Anima flowToastIn em 180ms e flowToastOut depois de 3s; posicao fixa no topo com safe-area."],
            ["Tour", "TourModal", "Overlay fade-in/fade-out 0.4s, caixa de highlight com transicao all 0.45s e tooltip reposicionado por targetRect."],
            ["Filtro candidatos", "SelectBase", "Chevron gira 180deg; dropdown abre/fecha com estado candidateFilterOpen."],
            ["Continuar flutuante", "SelectBase", "Shell fixo com transicao de opacity/transform/visibility; aparece ao atingir selecao minima ou subir scroll."],
            ["Cards", "StateCard/CandidateCard", "Hover/active usam transform translateY(-1px) ou scale(0.99) e box-shadow em 160ms."],
            ["Termometro", "CandidateCard + CSS", "Fill usa largura/progresso por variavel; estado featured/complete/tone muda cor e badge."],
            ["Compartilhamento", "ShareChoicePanel", "Modal fixo com backdrop blur, slides por transform/opacity/filter e botoes com active scale."],
            ["Navegacao", "BottomNavigation", "Indicador ativo usa transform scaleX e opacidade; drawers opcionais possuem optionsDrawerIn/optionsDesktopIn."],
        ],
        [2000, 2500, 4860],
    )

    doc.add_heading("9.1 Tons dos cards", level=2)
    add_table(
        doc,
        ["Tone/classe", "Regra"],
        [
            ["success", "Score positivo >= 7 e chance < 100; candidato interessante."],
            ["danger", "Score > 0 e < 7; abre modal de nota baixa antes de efetivar selecao."],
            ["new", "Sem score/candidato ingressante; pode usar nota partidaria quando disponivel."],
            ["neutral", "Candidato ja escolhido em outro passo ou chance >= 100."],
            ["visitor", "Campos personalizados bloqueados por falta de login."],
            ["fire-featured", "Candidato com melhor relacao nota/viabilidade ainda abaixo de 100%."],
            ["viability-complete", "Chance >= 100; modal alerta que pode nao precisar de mais votos."],
        ],
        [2400, 6960],
    )


def add_frontend_logic(doc: Document) -> None:
    doc.add_heading("10. Logicas de front-end", level=1)
    add_table(
        doc,
        ["Area", "Logica detalhada"],
        [
            ["Autenticacao", "useAuthState(auth) alimenta UserProvider. Sem user, limpa userData/elegibilidade/filtro local. Com user, cria/migra users/{uid}, assina snapshot do perfil e da elegibilidade."],
            ["Filtro ativo", "Persistido em localStorage com chave plano-voto:filtro-ativo e resetado no logout."],
            ["Busca", "Home e candidatos usam useDeferredValue para desacoplar digitacao da filtragem. normalizeSearch remove acentos e padroniza comparacao."],
            ["Estado do fluxo", "getBallotEstado/readVisitorBallotDraft determinam UF ativa antes de buscar candidatos; se nao houver UF, candidato redireciona para /home."],
            ["Fetch de candidatos", "Tenta cache local; renderiza stale-cache quando disponivel; busca rede se cache nao estiver fresco; busca tallies por estado e aplica ao modelo visual."],
            ["Ordenacao", "Candidatos bloqueados ficam depois; demais sao agrupados por destaque, boa nota/chance, chance 100, nota baixa e sem nota; desempate por chance, score e nome."],
            ["Selecao", "UI atualiza otimista, persiste etapa, aplica delta local nos tallies e tenta refresh remoto; em erro reverte selecionados e mostra modal."],
            ["Progresso", "getBallotProgress considera estado, deputado >=1 e senadores >=2; nextRoute volta para a etapa incompleta."],
            ["MeuPlano", "Combina snapshot local, draft remoto, fetchCandidatesByIds e fetchCandidateTallies para mostrar detalhes mais atuais sem perder fallback."],
            ["Debug", "flowLog/flowWarn/flowError registram eventos de auth, drafts, selecao, fetches e erros em desenvolvimento."],
        ],
        [1800, 7560],
    )

    doc.add_heading("10.1 Estados de erro tratados", level=2)
    add_bullets(
        doc,
        [
            "Firebase sem variaveis VITE_* usa localPreviewConfig e mostra aviso no console; login real fica indisponivel.",
            "Falha de fetch remoto de draft cai para rascunho local para manter navegacao.",
            "Falha de tallies nao impede lista; usa cache ou chance calculada com dados disponiveis.",
            "Falha de persistencia mostra ConfirmModal 'NAO FOI POSSIVEL SALVAR' e reverte selecao otimista.",
            "Troca de estado com escolhas existentes exige confirmacao para nao apagar rascunho sem consentimento.",
            "Resgate de QR expirado/usado mostra erro claro e botao de retorno.",
        ],
    )


def add_pwa_deploy_accessibility(doc: Document) -> None:
    doc.add_heading("11. PWA, deploy e acessibilidade", level=1)
    add_table(
        doc,
        ["Tema", "Detalhes"],
        [
            ["Manifest", "Nome nossovoto.org, display standalone/minimal-ui, orientacao portrait-primary, icones 192/512 maskable e tema branco."],
            ["Service worker", "Em producao, registra /sw.js no load. Cache APP shell e STATIC assets; navegacao usa timeout de 4500ms e fallback para index.html."],
            ["Atualizacao", "registerServiceWorker chama registration.update() a cada 1h para buscar versoes novas."],
            ["Deploy Vercel/Firebase", "Rewrites para index.html garantem SPA; headers cacheiam assets imutaveis e impedem cache de sw.js."],
            ["Bundle", "Build gerou aviso de chunk >450kB no index JS; o app ja usa lazy import nas paginas, mas o chunk base ainda pode ser refinado com manualChunks."],
            ["Privacidade", "Banner explica cookies tecnicos, Firebase/Google, localStorage, cache de candidatos e ausencia de camera/microfone/geolocalizacao."],
        ],
        [2100, 7260],
    )

    doc.add_heading("11.1 Acessibilidade implementada", level=2)
    add_bullets(
        doc,
        [
            "Loading e FlowToast usam role=status e aria-live=polite.",
            "Modais usam role=dialog, aria-modal=true e aria-labelledby com useId.",
            "Botoes iconicos possuem aria-label/title quando necessario.",
            "BottomNavigation usa aria-current=page, aria-disabled e disabled nas etapas indisponiveis.",
            "Campos de busca usam input type=search e labels visuais.",
            "Reset global define :focus-visible com outline consistente.",
            "prefers-reduced-motion reduz transicoes e animacoes para usuarios sensiveis a movimento.",
        ],
    )

    doc.add_heading("12. Checklist operacional", level=1)
    add_table(
        doc,
        ["Item", "Status observado"],
        [
            ["Build local", "npm run build executado com sucesso em 28/05/2026."],
            ["Rotas principais", "Sobre, Home, Escolher Deputado, Meu Plano e Login foram abertas no app local."],
            ["Capturas", "5 imagens salvas em docs/generated-app-docs/screenshots."],
            ["PWA", "Manifest e service worker configurados para producao."],
            ["Firestore", "Regras cobrem users, private choiceConfig, publicCandidateChoices, candidatos, partidos, elections e denies finais."],
            ["Pontos futuros", "Revisar tamanho do chunk base, anonimato forte de publicCandidateChoices e cobertura automatizada de UI."],
        ],
        [2600, 6760],
    )


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    package = json.loads((ROOT / "package.json").read_text(encoding="utf-8"))
    doc = Document()
    style_document(doc)
    add_title_page(doc, package)
    add_section_intro(doc)
    add_routing_and_boot(doc)
    add_data_logic(doc)
    add_screen_docs(doc)
    add_components(doc)
    add_responsiveness(doc)
    add_transitions(doc)
    add_frontend_logic(doc)
    add_pwa_deploy_accessibility(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
